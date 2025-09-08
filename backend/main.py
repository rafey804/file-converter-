import os
import shutil
import tempfile
import asyncio
from datetime import datetime, timedelta
from typing import List, Optional
import uuid
import zipfile
from contextlib import asynccontextmanager
 
from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, status
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials

# Try importing required packages, provide fallbacks if not available
try:
    import aiofiles
    AIOFILES_AVAILABLE = True
except ImportError:
    AIOFILES_AVAILABLE = False
    print("Warning: aiofiles not available. Install with: pip install aiofiles")

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    print("Warning: PyPDF2 not available. Install with: pip install PyPDF2")

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    print("Warning: reportlab not available. Install with: pip install reportlab")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Install with: pip install python-docx")

try:
    from pdf2image import convert_from_path, convert_from_bytes
    from PIL import Image
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    print("Warning: pdf2image not available. Install with: pip install pdf2image Pillow")

import io

# Rate limiting and security
from collections import defaultdict
import time

# Configuration
UPLOAD_DIR = "uploads"
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
ALLOWED_PDF_TYPES = ["application/pdf"]
ALLOWED_WORD_TYPES = [
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword"
]

# Rate limiting storage
request_counts = defaultdict(list)
RATE_LIMIT = 10  # requests per minute
RATE_WINDOW = 60  # seconds

# Security
security = HTTPBearer(auto_error=False)

# Ensure upload directory exists
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Function to find poppler path
def find_poppler_path():
    """Find poppler installation path on Windows"""
    possible_paths = [
        r"C:\poppler\Library\bin",
        r"C:\poppler-windows\Library\bin",
        r"C:\Program Files\poppler\bin",
        r"C:\Program Files (x86)\poppler\bin",
        r"C:\poppler\bin",
        r"C:\poppler-windows\bin"
    ]
    
    for path in possible_paths:
        if os.path.exists(path) and os.path.exists(os.path.join(path, "pdftoppm.exe")):
            print(f"Found poppler at: {path}")
            return path
    
    print("Poppler not found in common locations")
    return None

# Find poppler path at startup
POPPLER_PATH = find_poppler_path()

# Utility function to validate PDF file
def validate_pdf_file(file_path: str) -> bool:
    """Validate if the PDF file is readable"""
    try:
        if not PYPDF2_AVAILABLE:
            return True  # Skip validation if PyPDF2 not available
            
        with open(file_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            # Try to get page count
            page_count = len(pdf_reader.pages)
            print(f"PDF validation: {page_count} pages found")
            return page_count > 0
    except Exception as e:
        print(f"PDF validation failed: {e}")
        return False

# Enhanced PDF to Images converter with multiple fallback methods
async def pdf_to_images_converter(pdf_path: str, output_dir: str) -> List[str]:
    """Convert PDF pages to images with enhanced error handling and fallbacks"""
    if not PDF2IMAGE_AVAILABLE:
        print("PDF2IMAGE not available")
        return []
    
    try:
        print(f"Converting PDF: {pdf_path}")
        print(f"Output directory: {output_dir}")
        print(f"Using poppler path: {POPPLER_PATH}")
        
        # Validate PDF file first
        if not validate_pdf_file(pdf_path):
            print("PDF validation failed")
            return []
        
        # Method 1: Try with explicit poppler path and conservative settings
        images = None
        
        if POPPLER_PATH:
            try:
                print("Trying Method 1: With explicit poppler path")
                images = convert_from_path(
                    pdf_path, 
                    dpi=150,  # Lower DPI to reduce memory usage
                    poppler_path=POPPLER_PATH,
                    first_page=1,
                    last_page=None,
                    thread_count=1,  # Single thread to avoid issues
                    grayscale=False,
                    size=None,
                    transparent=False,
                    single_file=False,
                    output_folder=None,
                    output_file=None,
                    strict=False  # Less strict parsing
                )
                print(f"Method 1 successful: {len(images)} pages")
            except Exception as e:
                print(f"Method 1 failed: {e}")
                images = None
        
        # Method 2: Try without explicit poppler path
        if images is None:
            try:
                print("Trying Method 2: Without explicit poppler path")
                images = convert_from_path(
                    pdf_path, 
                    dpi=150,
                    thread_count=1,
                    strict=False
                )
                print(f"Method 2 successful: {len(images)} pages")
            except Exception as e:
                print(f"Method 2 failed: {e}")
                images = None
        
        # Method 3: Try with bytes (load file into memory)
        if images is None:
            try:
                print("Trying Method 3: Convert from bytes")
                with open(pdf_path, 'rb') as pdf_file:
                    pdf_bytes = pdf_file.read()
                
                if POPPLER_PATH:
                    images = convert_from_bytes(
                        pdf_bytes,
                        dpi=150,
                        poppler_path=POPPLER_PATH,
                        thread_count=1,
                        strict=False
                    )
                else:
                    images = convert_from_bytes(
                        pdf_bytes,
                        dpi=150,
                        thread_count=1,
                        strict=False
                    )
                print(f"Method 3 successful: {len(images)} pages")
            except Exception as e:
                print(f"Method 3 failed: {e}")
                images = None
        
        # Method 4: Try page by page conversion (fallback for problematic PDFs)
        if images is None:
            try:
                print("Trying Method 4: Page by page conversion")
                images = []
                page_num = 1
                
                while True:
                    try:
                        if POPPLER_PATH:
                            page_images = convert_from_path(
                                pdf_path,
                                dpi=150,
                                poppler_path=POPPLER_PATH,
                                first_page=page_num,
                                last_page=page_num,
                                strict=False
                            )
                        else:
                            page_images = convert_from_path(
                                pdf_path,
                                dpi=150,
                                first_page=page_num,
                                last_page=page_num,
                                strict=False
                            )
                        
                        if page_images:
                            images.extend(page_images)
                            print(f"Converted page {page_num}")
                            page_num += 1
                        else:
                            break
                            
                        # Safety limit to prevent infinite loop
                        if page_num > 1000:
                            print("Reached page limit (1000)")
                            break
                            
                    except Exception as e:
                        print(f"Failed to convert page {page_num}: {e}")
                        break
                
                if images:
                    print(f"Method 4 successful: {len(images)} pages")
            except Exception as e:
                print(f"Method 4 failed: {e}")
                images = None
        
        # If all methods failed
        if not images:
            print("All conversion methods failed")
            return []
        
        # Save images
        image_paths = []
        
        for i, image in enumerate(images):
            try:
                image_filename = f"page_{i+1:03d}.png"
                image_path = os.path.join(output_dir, image_filename)
                
                # Save image with high quality
                image.save(image_path, 'PNG', quality=95, optimize=True)
                image_paths.append(image_path)
                print(f"Saved page {i+1} as {image_filename}")
                
                # Close the image to free memory
                image.close()
                
            except Exception as e:
                print(f"Failed to save page {i+1}: {e}")
                continue
        
        return image_paths
        
    except Exception as e:
        print(f"PDF to images conversion error: {e}")
        import traceback
        traceback.print_exc()
        return []

# Lifespan event handler
@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    print("PDF Converter API started successfully!")
    print(f"Upload directory: {UPLOAD_DIR}")
    print("Available dependencies:")
    print(f"  - aiofiles: {AIOFILES_AVAILABLE}")
    print(f"  - PyPDF2: {PYPDF2_AVAILABLE}")
    print(f"  - reportlab: {REPORTLAB_AVAILABLE}")
    print(f"  - python-docx: {DOCX_AVAILABLE}")
    print(f"  - pdf2image: {PDF2IMAGE_AVAILABLE}")
    if POPPLER_PATH:
        print(f"  - poppler path: {POPPLER_PATH}")
    else:
        print("  - poppler: NOT FOUND")
    cleanup_old_files()
    yield
    # Shutdown
    print("Shutting down PDF Converter API...")

# Create FastAPI app with lifespan
app = FastAPI(
    title="PDF Converter API",
    description="Convert between PDF, Word and merge PDFs",
    version="1.0.0",
    lifespan=lifespan
)

# CORS Configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000","http://flipfilex.com"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Utility Functions
def cleanup_old_files():
    """Remove files older than 1 hour"""
    try:
        for filename in os.listdir(UPLOAD_DIR):
            filepath = os.path.join(UPLOAD_DIR, filename)
            if os.path.isfile(filepath):
                file_age = time.time() - os.path.getctime(filepath)
                if file_age > 3600:  # 1 hour
                    os.remove(filepath)
    except Exception as e:
        print(f"Cleanup error: {e}")

def check_rate_limit(client_ip: str):
    """Simple rate limiting"""
    current_time = time.time()
    
    # Clean old requests
    request_counts[client_ip] = [
        req_time for req_time in request_counts[client_ip]
        if current_time - req_time < RATE_WINDOW
    ]
    
    # Check rate limit
    if len(request_counts[client_ip]) >= RATE_LIMIT:
        raise HTTPException(
            status_code=429,
            detail="Rate limit exceeded. Try again later."
        )
    
    # Add current request
    request_counts[client_ip].append(current_time)

def validate_file_size(file: UploadFile):
    """Validate file size"""
    if hasattr(file, 'size') and file.size > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413,
            detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB"
        )

def generate_unique_filename(original_filename: str, extension: str = None) -> str:
    """Generate unique filename"""
    if extension:
        return f"{uuid.uuid4()}{extension}"
    else:
        name, ext = os.path.splitext(original_filename)
        return f"{uuid.uuid4()}{ext}"

# Conversion Functions
async def pdf_to_word_converter(pdf_path: str, output_path: str) -> bool:
    """Convert PDF to Word document"""
    if not PYPDF2_AVAILABLE or not DOCX_AVAILABLE:
        return False
    
    try:
        # Create new Word document
        doc = Document()
        doc.add_heading('Converted from PDF', 0)
        
        # Read PDF and extract text
        with open(pdf_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text.strip():
                    doc.add_heading(f'Page {page_num + 1}', level=1)
                    doc.add_paragraph(text)
                else:
                    doc.add_paragraph(f"[Page {page_num + 1} - No extractable text]")
        
        # Save Word document
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"PDF to Word conversion error: {e}")
        return False

async def word_to_pdf_converter(word_path: str, output_path: str) -> bool:
    """Convert Word document to PDF"""
    if not DOCX_AVAILABLE or not REPORTLAB_AVAILABLE:
        return False
    
    try:
        # Read Word document
        doc = Document(word_path)
        
        # Create PDF using reportlab
        pdf_doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        content = []
        
        # Extract text from Word document
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                p = Paragraph(paragraph.text, styles['Normal'])
                content.append(p)
                content.append(Spacer(1, 12))
        
        if not content:
            content.append(Paragraph("No content found in document", styles['Normal']))
        
        pdf_doc.build(content)
        return True
        
    except Exception as e:
        print(f"Word to PDF conversion error: {e}")
        return False

async def merge_pdfs(pdf_paths: List[str], output_path: str) -> bool:
    """Merge multiple PDFs into one"""
    if not PYPDF2_AVAILABLE:
        return False
    
    try:
        pdf_merger = PyPDF2.PdfMerger()
        
        for pdf_path in pdf_paths:
            with open(pdf_path, 'rb') as pdf_file:
                pdf_merger.append(pdf_file)
        
        with open(output_path, 'wb') as output_file:
            pdf_merger.write(output_file)
        
        pdf_merger.close()
        return True
        
    except Exception as e:
        print(f"PDF merge error: {e}")
        return False

# File I/O helper functions
async def write_file(file_path: str, content: bytes):
    """Write file with fallback"""
    if AIOFILES_AVAILABLE:
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(content)
    else:
        with open(file_path, 'wb') as f:
            f.write(content)

# API Routes
@app.get("/")
async def root():
    return {"message": "PDF Converter API is running"}

@app.get("/health")
async def health_check():
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "version": "1.0.0",
        "dependencies": {
            "aiofiles": AIOFILES_AVAILABLE,
            "PyPDF2": PYPDF2_AVAILABLE,
            "reportlab": REPORTLAB_AVAILABLE,
            "python-docx": DOCX_AVAILABLE,
            "pdf2image": PDF2IMAGE_AVAILABLE and (POPPLER_PATH is not None)
        }
    }

@app.post("/convert/pdf-to-word")
async def convert_pdf_to_word(
    file: UploadFile = File(...),
    credentials: Optional[HTTPAuthorizationCredentials] = Depends(security)
):
    if not PYPDF2_AVAILABLE or not DOCX_AVAILABLE:
        raise HTTPException(
            status_code=503, 
            detail="PDF to Word conversion not available. Missing dependencies: PyPDF2 or python-docx"
        )
    
    # Rate limiting
    client_ip = "127.0.0.1"
    check_rate_limit(client_ip)
    cleanup_old_files()
    
    # Validate file
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are allowed")
    
    validate_file_size(file)
    
    try:
        # Save uploaded file
        input_filename = generate_unique_filename(file.filename)
        input_path = os.path.join(UPLOAD_DIR, input_filename)
        
        content = await file.read()
        await write_file(input_path, content)
        
        # Generate output filename
        output_filename = generate_unique_filename(file.filename, '.docx')
        output_path = os.path.join(UPLOAD_DIR, output_filename)
        
        # Convert PDF to Word
        success = await pdf_to_word_converter(input_path, output_path)
        
        if not success:
            if os.path.exists(input_path):
                os.remove(input_path)
            raise HTTPException(status_code=500, detail="Conversion failed")
        
        # Cleanup input file
        if os.path.exists(input_path):
            os.remove(input_path)
        
        return {
            "message": "PDF converted to Word successfully",
            "download_url": f"/download/{output_filename}",
            "filename": output_filename
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion error: {str(e)}")

@app.post("/convert/word-to-pdf")
async def convert_word_to_pdf(
    file: UploadFile = File(...),
    credentials: Optional[HTTPAuthorizationCredentials] = Depends(security)
):
    if not DOCX_AVAILABLE or not REPORTLAB_AVAILABLE:
        raise HTTPException(
            status_code=503, 
            detail="Word to PDF conversion not available. Missing dependencies: python-docx or reportlab"
        )
    
    client_ip = "127.0.0.1"
    check_rate_limit(client_ip)
    cleanup_old_files()
    
    # Validate file
    if not (file.filename.lower().endswith('.docx') or file.filename.lower().endswith('.doc')):
        raise HTTPException(status_code=400, detail="Only Word documents (.docx, .doc) are allowed")
    
    validate_file_size(file)
    
    try:
        # Save uploaded file
        input_filename = generate_unique_filename(file.filename)
        input_path = os.path.join(UPLOAD_DIR, input_filename)
        
        content = await file.read()
        await write_file(input_path, content)
        
        # Generate output filename
        output_filename = generate_unique_filename(file.filename, '.pdf')
        output_path = os.path.join(UPLOAD_DIR, output_filename)
        
        # Convert Word to PDF
        success = await word_to_pdf_converter(input_path, output_path)
        
        if not success:
            if os.path.exists(input_path):
                os.remove(input_path)
            raise HTTPException(status_code=500, detail="Conversion failed")
        
        # Cleanup input file
        if os.path.exists(input_path):
            os.remove(input_path)
        
        return {
            "message": "Word document converted to PDF successfully",
            "download_url": f"/download/{output_filename}",
            "filename": output_filename
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion error: {str(e)}")

@app.post("/convert/merge-pdf")
async def merge_pdf_files(
    files: List[UploadFile] = File(...),
    credentials: Optional[HTTPAuthorizationCredentials] = Depends(security)
):
    if not PYPDF2_AVAILABLE:
        raise HTTPException(
            status_code=503, 
            detail="PDF merge not available. Missing dependency: PyPDF2"
        )
    
    client_ip = "127.0.0.1"
    check_rate_limit(client_ip)
    cleanup_old_files()
    
    if len(files) < 2:
        raise HTTPException(status_code=400, detail="At least 2 PDF files are required for merging")
    
    if len(files) > 10:
        raise HTTPException(status_code=400, detail="Maximum 10 files allowed for merging")
    
    input_paths = []
    
    try:
        # Save all uploaded files
        for file in files:
            if not file.filename.lower().endswith('.pdf'):
                raise HTTPException(status_code=400, detail="Only PDF files are allowed")
            
            validate_file_size(file)
            
            input_filename = generate_unique_filename(file.filename)
            input_path = os.path.join(UPLOAD_DIR, input_filename)
            input_paths.append(input_path)
            
            content = await file.read()
            await write_file(input_path, content)
        
        # Generate output filename
        output_filename = generate_unique_filename("merged_document.pdf", '.pdf')
        output_path = os.path.join(UPLOAD_DIR, output_filename)
        
        # Merge PDFs
        success = await merge_pdfs(input_paths, output_path)
        
        if not success:
            raise HTTPException(status_code=500, detail="PDF merge failed")
        
        # Cleanup input files
        for input_path in input_paths:
            if os.path.exists(input_path):
                os.remove(input_path)
        
        return {
            "message": f"{len(files)} PDF files merged successfully",
            "download_url": f"/download/{output_filename}",
            "filename": output_filename
        }
        
    except Exception as e:
        # Cleanup on error
        for input_path in input_paths:
            if os.path.exists(input_path):
                os.remove(input_path)
        raise HTTPException(status_code=500, detail=f"Merge error: {str(e)}")

@app.post("/convert/pdf-to-images")
async def convert_pdf_to_images(
    file: UploadFile = File(...),
    credentials: Optional[HTTPAuthorizationCredentials] = Depends(security)
):
    if not PDF2IMAGE_AVAILABLE:
        raise HTTPException(
            status_code=503, 
            detail="PDF to Images conversion not available. Missing dependencies: pdf2image or Pillow"
        )
    
    # Rate limiting
    client_ip = "127.0.0.1"
    check_rate_limit(client_ip)
    cleanup_old_files()
    
    # Validate file
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are allowed")
    
    validate_file_size(file)
    
    input_path = None
    output_dir = None
    
    try:
        # Save uploaded file
        input_filename = generate_unique_filename(file.filename)
        input_path = os.path.join(UPLOAD_DIR, input_filename)
        
        content = await file.read()
        await write_file(input_path, content)
        
        print(f"Saved uploaded file to: {input_path}")
        
        # Create output directory for images
        output_dir = os.path.join(UPLOAD_DIR, f"images_{uuid.uuid4()}")
        os.makedirs(output_dir, exist_ok=True)
        
        print(f"Created output directory: {output_dir}")
        
        # Convert PDF to images with enhanced error handling
        image_paths = await pdf_to_images_converter(input_path, output_dir)
        
        if not image_paths:
            raise HTTPException(status_code=500, detail="PDF conversion failed. Could not extract images from PDF. The PDF file might be corrupted, password-protected, or in an unsupported format.")
        
        print(f"Successfully converted to {len(image_paths)} images")
        
        # Create ZIP file with all images
        zip_filename = f"pdf_images_{uuid.uuid4().hex[:8]}.zip"
        zip_path = os.path.join(UPLOAD_DIR, zip_filename)
        
        print(f"Creating ZIP file: {zip_path}")
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for image_path in image_paths:
                # Add file to zip with just the filename (not full path)
                zipf.write(image_path, os.path.basename(image_path))
                print(f"Added {os.path.basename(image_path)} to ZIP")
        
        # Cleanup input file and images directory
        if os.path.exists(input_path):
            os.remove(input_path)
        if os.path.exists(output_dir):
            shutil.rmtree(output_dir)
        
        print(f"Conversion completed successfully. ZIP file: {zip_filename}")
        
        return {
            "message": f"PDF converted to {len(image_paths)} images successfully",
            "download_url": f"/download/{zip_filename}",
            "filename": zip_filename,
            "image_count": len(image_paths)
        }
        
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        print(f"Error in pdf-to-images conversion: {e}")
        import traceback
        traceback.print_exc()
        
        # Cleanup on error
        if input_path and os.path.exists(input_path):
            os.remove(input_path)
        if output_dir and os.path.exists(output_dir):
            shutil.rmtree(output_dir)
        raise HTTPException(status_code=500, detail=f"Conversion error: {str(e)}")

@app.get("/download/{filename}")
async def download_file(filename: str):
    file_path = os.path.join(UPLOAD_DIR, filename)
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    # Security check
    if ".." in filename or "/" in filename or "\\" in filename:
        raise HTTPException(status_code=400, detail="Invalid filename")
    
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type='application/octet-stream'
    )

# Error handlers
@app.exception_handler(413)
async def file_too_large_handler(request, exc):
    return JSONResponse(
        status_code=413,
        content={"detail": "File too large"}
    )

@app.exception_handler(500)
async def internal_error_handler(request, exc):
    return JSONResponse(
        status_code=500,
        content={"detail": "Internal server error"}
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=5000, reload=True)